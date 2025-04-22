from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_documentation():
    doc = Document()

    # Sarlavha
    title = doc.add_heading('Movie API Dokumentatsiyasi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Ushbu hujjat API’ning to‘liq tavsifi bo‘lib, frontend ishlab chiquvchilar uchun mo‘ljallangan. '
                      'Unda modellar, serializerlar, ruxsatlar, admin paneli, JWT autentifikatsiyasi va Swagger integratsiyasi keltirilgan.')

    # 1. Modellar
    doc.add_heading('1. Modellar', level=1)
    doc.add_paragraph('Quyida loyihadagi barcha modellar va ularning maydonlari keltiriladi.')

    doc.add_heading('1.1 User', level=2)
    doc.add_paragraph('Foydalanuvchi modeli, Django AbstractUser’dan meros oladi.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('username: Foydalanuvchi nomi (CharField).', style='List Bullet')
    doc.add_paragraph('email: Elektron pochta (EmailField).', style='List Bullet')
    doc.add_paragraph('password: Parol (hash qilingan).', style='List Bullet')
    doc.add_paragraph('first_name: Ism (CharField).', style='List Bullet')
    doc.add_paragraph('last_name: Familiya (CharField).', style='List Bullet')
    doc.add_paragraph('role: Rol (CharField, choices: admin yoki exploiter).', style='List Bullet')

    doc.add_heading('1.2 Admin', level=2)
    doc.add_paragraph('Admin foydalanuvchisi uchun profil.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('user: User modeli bilan OneToOneField aloqa (related_name="admin_profile").',
                      style='List Bullet')
    doc.add_paragraph('phone_number: Telefon raqami (CharField, max_length=13).', style='List Bullet')
    doc.add_paragraph('addres: Manzil (CharField, max_length=255, tavsiya: address sifatida tuzatish).',
                      style='List Bullet')

    doc.add_heading('1.3 Exploiter', level=2)
    doc.add_paragraph('Exploiter foydalanuvchisi uchun profil.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('user: User modeli bilan OneToOneField aloqa (related_name="exploiter_profile").',
                      style='List Bullet')
    doc.add_paragraph('phone_number: Telefon raqami (CharField, max_length=13).', style='List Bullet')

    doc.add_heading('1.4 Genre', level=2)
    doc.add_paragraph('Kino janri modeli.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('name: Janr nomi (CharField, max_length=100).', style='List Bullet')

    doc.add_heading('1.5 Company', level=2)
    doc.add_paragraph('Kino ishlab chiqaruvchi kompaniya modeli.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('name: Kompaniya nomi (CharField, max_length=100).', style='List Bullet')
    doc.add_paragraph('description: Tavsif (TextField, ixtiyoriy).', style='List Bullet')

    doc.add_heading('1.6 Actor', level=2)
    doc.add_paragraph('Aktyor modeli.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('first_name: Ism (CharField, max_length=100).', style='List Bullet')
    doc.add_paragraph('last_name: Familiya (CharField, max_length=100).', style='List Bullet')
    doc.add_paragraph('birth_date: Tug‘ilgan sana (DateField, ixtiyoriy).', style='List Bullet')
    doc.add_paragraph('photo: Rasm (ImageField, ixtiyoriy).', style='List Bullet')

    doc.add_heading('1.7 Film', level=2)
    doc.add_paragraph('Kino modeli.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('name: Kino nomi (CharField, max_length=150).', style='List Bullet')
    doc.add_paragraph('views: Ko‘rishlar soni (PositiveIntegerField, default=0).', style='List Bullet')
    doc.add_paragraph('genre: Janr bilan ForeignKey aloqa (SET_NULL).', style='List Bullet')
    doc.add_paragraph('company: Kompaniya bilan ForeignKey aloqa (SET_NULL).', style='List Bullet')
    doc.add_paragraph('actors: Aktyorlar bilan ManyToManyField aloqa.', style='List Bullet')

    doc.add_heading('1.8 Comment', level=2)
    doc.add_paragraph('Foydalanuvchi izohlari modeli.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('user: Foydalanuvchi bilan ForeignKey aloqa.', style='List Bullet')
    doc.add_paragraph('film: Kino bilan ForeignKey aloqa.', style='List Bullet')
    doc.add_paragraph('text: Izoh matni (TextField).', style='List Bullet')
    doc.add_paragraph('created_at: Yaratilgan vaqt (DateTimeField, auto_now_add=True).', style='List Bullet')

    doc.add_heading('1.9 Rating', level=2)
    doc.add_paragraph('Kino baholash modeli.')
    doc.add_paragraph('Maydonlar:', style='List Bullet')
    doc.add_paragraph('user: Foydalanuvchi bilan ForeignKey aloqa.', style='List Bullet')
    doc.add_paragraph('film: Kino bilan ForeignKey aloqa.', style='List Bullet')
    doc.add_paragraph('score: Baho (PositiveSmallIntegerField, 1-10).', style='List Bullet')
    doc.add_paragraph('created_at: Yaratilgan vaqt (DateTimeField, auto_now_add=True).', style='List Bullet')
    doc.add_paragraph('Meta: user va film birikmasi unik (unique_together).', style='List Bullet')

    # 2. Serializerlar
    doc.add_heading('2. Serializerlar', level=1)
    doc.add_paragraph('Serializerlar ma’lumotlarni JSON formatiga o‘tkazish va validatsiya qilish uchun ishlatiladi. '
                      'Quyida har bir serializerning maqsadi va logikasi, ayniqsa yangilangan funksionallik keltiriladi.')

    doc.add_heading('2.1 UserSerializer', level=2)
    doc.add_paragraph('Foydalanuvchi ma’lumotlarini serializatsiya qiladi.')
    doc.add_paragraph('Maydonlar: id, username, email, password (write_only), first_name, last_name, role')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('create_user metodi orqali yangi foydalanuvchi yaratiladi.', style='List Bullet')
    doc.add_paragraph('Parol hash qilinadi.', style='List Bullet')
    doc.add_paragraph('role maydoni default sifatida "exploiter" qabul qilinadi.', style='List Bullet')
    doc.add_paragraph('Update logikasi:', style='List Bullet')
    doc.add_paragraph('Mavjud foydalanuvchi ma’lumotlari yangilanadi.', style='List Bullet')
    doc.add_paragraph('Agar parol berilgan bo‘lsa, set_password orqali hash qilinadi.', style='List Bullet')

    doc.add_heading('2.2 AdminSerializer', level=2)
    doc.add_paragraph('Admin profilini serializatsiya qiladi, user maydoni nested UserSerializer orqali.')
    doc.add_paragraph('Maydonlar: id, user, phone_number, addres')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('user ma’lumotlari UserSerializer orqali yangi User yaratadi (role="admin").',
                      style='List Bullet')
    doc.add_paragraph('Admin obyekti yaratiladi va user bilan bog‘lanadi.', style='List Bullet')
    doc.add_paragraph('Update logikasi:', style='List Bullet')
    doc.add_paragraph('user ma’lumotlari UserSerializer orqali yangilanadi (role="admin" saqlanadi).',
                      style='List Bullet')
    doc.add_paragraph('phone_number va addres yangilanadi.', style='List Bullet')

    doc.add_heading('2.3 ExploiterSerializer', level=2)
    doc.add_paragraph('Exploiter profilini serializatsiya qiladi, user maydoni nested UserSerializer orqali.')
    doc.add_paragraph('Maydonlar: id, user, phone_number')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('user ma’lumotlari UserSerializer orqali yangi User yaratadi (role="exploiter").',
                      style='List Bullet')
    doc.add_paragraph('Exploiter obyekti yaratiladi va user bilan bog‘lanadi.', style='List Bullet')
    doc.add_paragraph('Update logikasi:', style='List Bullet')
    doc.add_paragraph('user ma’lumotlari UserSerializer orqali yangilanadi (role="exploiter" saqlanadi).',
                      style='List Bullet')
    doc.add_paragraph('phone_number yangilanadi.', style='List Bullet')

    doc.add_heading('2.4 GenreSerializer', level=2)
    doc.add_paragraph('Kino janrini serializatsiya qiladi. Yangilanish: yangi janr yaratish yoki mavjudini ishlatish.')
    doc.add_paragraph('Maydonlar: id, name')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('Agar berilgan nomdagi janr mavjud bo‘lsa, uni qaytaradi (get_or_create).', style='List Bullet')
    doc.add_paragraph('Aks holda, yangi janr yaratadi.', style='List Bullet')

    doc.add_heading('2.5 CompanySerializer', level=2)
    doc.add_paragraph(
        'Kompaniya ma’lumotlarini serializatsiya qiladi. Yangilanish: yangi kompaniya yaratish yoki mavjudini ishlatish.')
    doc.add_paragraph('Maydonlar: id, name, description')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('Agar berilgan nomdagi kompaniya mavjud bo‘lsa, uni qaytaradi (get_or_create).',
                      style='List Bullet')
    doc.add_paragraph('Aks holda, yangi kompaniya yaratadi, description ixtiyoriy.', style='List Bullet')

    doc.add_heading('2.6 ActorSerializer', level=2)
    doc.add_paragraph(
        'Aktyor ma’lumotlarini serializatsiya qiladi. Yangilanish: yangi aktyor yaratish yoki mavjudini ishlatish.')
    doc.add_paragraph('Maydonlar: id, first_name, last_name, birth_date, photo')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('Agar first_name va last_name bo‘yicha aktyor mavjud bo‘lsa, uni qaytaradi (get_or_create).',
                      style='List Bullet')
    doc.add_paragraph('Aks holda, yangi aktyor yaratadi, birth_date va photo ixtiyoriy.', style='List Bullet')

    doc.add_heading('2.7 FilmSerializer', level=2)
    doc.add_paragraph(
        'Kino ma’lumotlarini serializatsiya qiladi. Yangilanish: genre, company, actors uchun nested serializerlar va yangi/mavjud obyektlar bilan ishlash.')
    doc.add_paragraph('Maydonlar: id, name, views, genre, company, actors')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('Agar genre uchun JSON ma’lumot kelsa (masalan, {"name": "Sci-Fi"}), yangi Genre yaratiladi.',
                      style='List Bullet')
    doc.add_paragraph('Agar genre_id kelsa, mavjud janr ishlatiladi.', style='List Bullet')
    doc.add_paragraph('company va actors uchun xuddi shunday: JSON ma’lumot yoki company_id/actor_ids ishlatiladi.',
                      style='List Bullet')
    doc.add_paragraph('Film obyekti yaratiladi, actors ManyToManyField orqali qo‘shiladi.', style='List Bullet')
    doc.add_paragraph('Update logikasi:', style='List Bullet')
    doc.add_paragraph('Mavjud film yangilanadi, genre, company, actors uchun yangi ma’lumot yoki ID’lar ishlatiladi.',
                      style='List Bullet')
    doc.add_paragraph('Agar ma’lumot kelmagan bo‘lsa, eski qiymatlar saqlanadi.', style='List Bullet')

    doc.add_heading('2.8 CommentSerializer', level=2)
    doc.add_paragraph('Izoh ma’lumotlarini serializatsiya qiladi.')
    doc.add_paragraph('Maydonlar: id, user (read_only), film (read_only), text, created_at')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('user joriy foydalanuvchidan (request.user) olinadi.', style='List Bullet')
    doc.add_paragraph('film_id request.data dan olinadi.', style='List Bullet')
    doc.add_paragraph('Comment obyekti yaratiladi.', style='List Bullet')

    doc.add_heading('2.9 RatingSerializer', level=2)
    doc.add_paragraph('Baho ma’lumotlarini serializatsiya qiladi.')
    doc.add_paragraph('Maydonlar: id, user (read_only), film (read_only), score, created_at')
    doc.add_paragraph('Create logikasi:', style='List Bullet')
    doc.add_paragraph('user joriy foydalanuvchidan (request.user) olinadi.', style='List Bullet')
    doc.add_paragraph('film_id request.data dan olinadi.', style='List Bullet')
    doc.add_paragraph('Rating obyekti yaratiladi.', style='List Bullet')
    doc.add_paragraph('Update logikasi:', style='List Bullet')
    doc.add_paragraph('Faqat score maydoni yangilanadi.', style='List Bullet')

    # 3. Maxsus Ruxsatlar
    doc.add_heading('3. Maxsus Ruxsatlar', level=1)
    doc.add_paragraph(
        'Standart ruxsatlar (IsAdminUser, IsAuthenticated, AllowAny) o‘rniga permissions.py faylida maxsus ruxsat sinflari yaratildi.')

    doc.add_heading('3.1 IsAdmin', level=2)
    doc.add_paragraph('Faqat role="admin" foydalanuvchilar uchun ruxsat beradi.')
    doc.add_paragraph('Qo‘llanilishi: User va Admin ViewSet’larida.', style='List Bullet')

    doc.add_heading('3.2 IsAuthenticatedOrExploiter', level=2)
    doc.add_paragraph(
        'Autentifikatsiya qilingan foydalanuvchilar uchun ochiq, lekin exploiter faqat o‘z profilini ko‘radi.')
    doc.add_paragraph('Logika:', style='List Bullet')
    doc.add_paragraph(
        'has_object_permission orqali exploiter faqat o‘z user obyekti bilan bog‘langan profilni ko‘radi.',
        style='List Bullet')
    doc.add_paragraph('Adminlar barcha exploter profillarini ko‘rishi mumkin.', style='List Bullet')
    doc.add_paragraph('Qo‘llanilishi: ExploiterViewSet.', style='List Bullet')

    doc.add_heading('3.3 AllowReadOnlyForNonAdmins', level=2)
    doc.add_paragraph('GET so‘rovlari uchun hamma uchun ochiq, yozish (POST, PUT, DELETE) faqat adminlar uchun.')
    doc.add_paragraph('Qo‘llanilishi: Genre, Company, Actor, Film ViewSet’larida.', style='List Bullet')

    doc.add_heading('3.4 IsAuthenticatedForComments', level=2)
    doc.add_paragraph(
        'Autentifikatsiya qilingan foydalanuvchilar izoh yozishi mumkin, lekin faqat o‘z izohlarini ko‘radi.')
    doc.add_paragraph('Logika:', style='List Bullet')
    doc.add_paragraph('has_object_permission orqali foydalanuvchi faqat o‘z izohlarini ko‘radi.', style='List Bullet')
    doc.add_paragraph('Adminlar barcha izohlarni ko‘rishi mumkin.', style='List Bullet')
    doc.add_paragraph('Qo‘llanilishi: CommentViewSet.', style='List Bullet')

    doc.add_heading('3.5 IsAuthenticatedForRatings', level=2)
    doc.add_paragraph(
        'Autentifikatsiya qilingan foydalanuvchilar baho qo‘yishi mumkin, lekin faqat o‘z baholarini ko‘radi.')
    doc.add_paragraph('Logika:', style='List Bullet')
    doc.add_paragraph('has_object_permission orqali foydalanuvchi faqat o‘z baholarini ko‘radi.', style='List Bullet')
    doc.add_paragraph('Adminlar barcha baholarni ko‘rishi mumkin.', style='List Bullet')
    doc.add_paragraph('Qo‘llanilishi: RatingViewSet.', style='List Bullet')

    # 4. ViewSet’lar va URL’lar
    doc.add_heading('4. ViewSet’lar va URL’lar', level=1)
    doc.add_paragraph('Har bir model uchun ViewSet va unga mos keluvchi API endpointlari.')

    doc.add_heading('4.1 UserViewSet', level=2)
    doc.add_paragraph('Foydalanuvchilarni boshqarish uchun.')
    doc.add_paragraph('Ruxsatlar: IsAdmin (faqat adminlar).')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/users/')

    doc.add_heading('4.2 AdminViewSet', level=2)
    doc.add_paragraph('Admin profillarini boshqarish uchun.')
    doc.add_paragraph('Ruxsatlar: IsAdmin (faqat adminlar).')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/admins/')

    doc.add_heading('4.3 ExploiterViewSet', level=2)
    doc.add_paragraph('Exploiter profillarini boshqarish uchun.')
    doc.add_paragraph(
        'Ruxsatlar: IsAuthenticatedOrExploiter (autentifikatsiya qilingan foydalanuvchilar, faqat o‘z profilini ko‘radi).')
    doc.add_paragraph('Logika: Oddiy foydalanuvchilar faqat o‘z profilini ko‘radi, adminlar hammasini.')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/exploiters/')

    doc.add_heading('4.4 GenreViewSet', level=2)
    doc.add_paragraph('Janrlarni boshqarish uchun.')
    doc.add_paragraph('Ruxsatlar: AllowReadOnlyForNonAdmins (hamma ko‘rishi mumkin, faqat adminlar yozishi mumkin).')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/genres/')

    doc.add_heading('4.5 CompanyViewSet', level=2)
    doc.add_paragraph('Kompaniyalarni boshqarish uchun.')
    doc.add_paragraph('Ruxsatlar: AllowReadOnlyForNonAdmins.')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/companies/')

    doc.add_heading('4.6 ActorViewSet', level=2)
    doc.add_paragraph('Aktyorlarni boshqarish uchun.')
    doc.add_paragraph('Ruxsatlar: AllowReadOnlyForNonAdmins.')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/actors/')

    doc.add_heading('4.7 FilmViewSet', level=2)
    doc.add_paragraph('Kinolar boshqarish uchun.')
    doc.add_paragraph('Ruxsatlar: AllowReadOnlyForNonAdmins.')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/films/')

    doc.add_heading('4.8 CommentViewSet', level=2)
    doc.add_paragraph('Izohlarni boshqarish uchun.')
    doc.add_paragraph(
        'Ruxsatlar: IsAuthenticatedForComments (autentifikatsiya qilingan foydalanuvchilar, faqat o‘z izohlarini ko‘radi).')
    doc.add_paragraph('Logika: Foydalanuvchilar faqat o‘z izohlarini ko‘radi, adminlar hammasini.')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/comments/')

    doc.add_heading('4.9 RatingViewSet', level=2)
    doc.add_paragraph('Baholarni boshqarish uchun.')
    doc.add_paragraph(
        'Ruxsatlar: IsAuthenticatedForRatings (autentifikatsiya qilingan foydalanuvchilar, faqat o‘z baholarini ko‘radi).')
    doc.add_paragraph('Logika: Foydalanuvchilar faqat o‘z baholarini ko‘radi, adminlar hammasini.')
    doc.add_paragraph('URL: http://127.0.0.1:8000/app/v1/ratings/')

    # 5. Admin Paneli
    doc.add_heading('5. Admin Paneli', level=1)
    doc.add_paragraph(
        'Admin paneli har bir model uchun chiroyli va qulay interfeys bilan sozlandi. Har bir model uchun ModelAdmin sinflari yaratildi.')

    doc.add_heading('5.1 UserAdmin', level=2)
    doc.add_paragraph('Foydalanuvchilarni boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: username, email, role, first_name, last_name, is_staff.', style='List Bullet')
    doc.add_paragraph('list_filter: role, is_staff, is_active.', style='List Bullet')
    doc.add_paragraph('search_fields: username, email, first_name, last_name.', style='List Bullet')
    doc.add_paragraph('fields: username, email, password, first_name, last_name, role, is_staff, is_active.',
                      style='List Bullet')

    doc.add_heading('5.2 AdminAdmin', level=2)
    doc.add_paragraph('Admin profillarini boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: user, phone_number, addres.', style='List Bullet')
    doc.add_paragraph('list_filter: phone_number.', style='List Bullet')
    doc.add_paragraph('search_fields: user__username, phone_number, addres.', style='List Bullet')
    doc.add_paragraph('raw_id_fields: user.', style='List Bullet')

    doc.add_heading('5.3 ExploiterAdmin', level=2)
    doc.add_paragraph('Exploiter profillarini boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: user, phone_number.', style='List Bullet')
    doc.add_paragraph('list_filter: phone_number.', style='List Bullet')
    doc.add_paragraph('search_fields: user__username, phone_number.', style='List Bullet')
    doc.add_paragraph('raw_id_fields: user.', style='List Bullet')

    doc.add_heading('5.4 GenreAdmin', level=2)
    doc.add_paragraph('Janrlarni boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: name.', style='List Bullet')
    doc.add_paragraph('search_fields: name.', style='List Bullet')
    doc.add_paragraph('fields: name.', style='List Bullet')

    doc.add_heading('5.5 CompanyAdmin', level=2)
    doc.add_paragraph('Kompaniyalarni boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: name, description.', style='List Bullet')
    doc.add_paragraph('search_fields: name, description.', style='List Bullet')
    doc.add_paragraph('fields: name, description.', style='List Bullet')

    doc.add_heading('5.6 ActorAdmin', level=2)
    doc.add_paragraph('Aktyorlarni boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: first_name, last_name, birth_date.', style='List Bullet')
    doc.add_paragraph('list_filter: birth_date.', style='List Bullet')
    doc.add_paragraph('search_fields: first_name, last_name.', style='List Bullet')
    doc.add_paragraph('fields: first_name, last_name, birth_date, photo.', style='List Bullet')

    doc.add_heading('5.7 FilmAdmin', level=2)
    doc.add_paragraph('Kinolar boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: name, genre, company, views.', style='List Bullet')
    doc.add_paragraph('list_filter: genre, company.', style='List Bullet')
    doc.add_paragraph('search_fields: name.', style='List Bullet')
    doc.add_paragraph('fields: name, genre, company, actors, views.', style='List Bullet')
    doc.add_paragraph('filter_horizontal: actors (ManyToManyField uchun qulay interfeys).', style='List Bullet')
    doc.add_paragraph('raw_id_fields: genre, company.', style='List Bullet')

    doc.add_heading('5.8 CommentAdmin', level=2)
    doc.add_paragraph('Izohlarni boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: user, film, text_preview, created_at.', style='List Bullet')
    doc.add_paragraph('list_filter: created_at, user, film.', style='List Bullet')
    doc.add_paragraph('search_fields: user__username, film__name, text.', style='List Bullet')
    doc.add_paragraph('fields: user, film, text, created_at.', style='List Bullet')
    doc.add_paragraph('raw_id_fields: user, film.', style='List Bullet')
    doc.add_paragraph('readonly_fields: created_at.', style='List Bullet')
    doc.add_paragraph('text_preview: Izoh matnining qisqa ko‘rinishi.', style='List Bullet')

    doc.add_heading('5.9 RatingAdmin', level=2)
    doc.add_paragraph('Baholarni boshqarish uchun.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('list_display: user, film, score, created_at.', style='List Bullet')
    doc.add_paragraph('list_filter: score, created_at, user, film.', style='List Bullet')
    doc.add_paragraph('search_fields: user__username, film__name.', style='List Bullet')
    doc.add_paragraph('fields: user, film, score, created_at.', style='List Bullet')
    doc.add_paragraph('raw_id_fields: user, film.', style='List Bullet')
    doc.add_paragraph('readonly_fields: created_at.', style='List Bullet')

    # 6. JWT Autentifikatsiyasi
    doc.add_heading('6. JWT Autentifikatsiyasi', level=1)
    doc.add_paragraph('API’da autentifikatsiya uchun djangorestframework-simplejwt kutubxonasi ishlatildi. '
                      'Bu foydalanuvchilarga access va refresh tokenlari orqali kirish imkonini beradi.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('settings.py’da REST_FRAMEWORK va SIMPLE_JWT sozlamalari qo‘shildi.', style='List Bullet')
    doc.add_paragraph('Access token muddati: 60 daqiqa.', style='List Bullet')
    doc.add_paragraph('Refresh token muddati: 1 kun.', style='List Bullet')
    doc.add_paragraph('Token olish: POST /api/token/ (username va password bilan).', style='List Bullet')
    doc.add_paragraph('Token yangilash: POST /api/token/refresh/ (refresh token bilan).', style='List Bullet')
    doc.add_paragraph('Swagger’da token kiritish: Bearer <access_token> formatida.', style='List Bullet')

    # 7. Swagger Integratsiyasi
    doc.add_heading('7. Swagger Integratsiyasi', level=1)
    doc.add_paragraph(
        'API dokumentatsiyasi uchun drf-yasg kutubxonasi ishlatildi, bu Swagger UI va ReDoc interfeyslarini taqdim etadi.')
    doc.add_paragraph('Sozlamalar:', style='List Bullet')
    doc.add_paragraph('settings.py’da drf_yasg INSTALLED_APPS’ga qo‘shildi.', style='List Bullet')
    doc.add_paragraph('urls.py’da /swagger/ va /redoc/ endpoint’lari qo‘shildi.', style='List Bullet')
    doc.add_paragraph('Swagger UI: http://127.0.0.1:8000/swagger/', style='List Bullet')
    doc.add_paragraph('ReDoc: http://127.0.0.1:8000/redoc/', style='List Bullet')
    doc.add_paragraph('JWT integratsiyasi: SWAGGER_SETTINGS’da Bearer token sozlamasi qo‘shildi.', style='List Bullet')
    doc.add_paragraph('Foydalanish: “Authorize” tugmasi orqali Bearer <access_token> kiritiladi.', style='List Bullet')

    # 8. Foydalanuvchilar
    doc.add_heading('8. Foydalanuvchilar (User, Admin, Exploiter)', level=1)
    doc.add_paragraph('Foydalanuvchilar bilan ishlashda quyidagi ishlar amalga oshirildi:')
    doc.add_paragraph('User modeli:', style='List Bullet')
    doc.add_paragraph('AbstractUser’dan meros oladi, role maydoni qo‘shildi (admin yoki exploiter).',
                      style='List Bullet')
    doc.add_paragraph('JWT autentifikatsiyasi orqali foydalanuvchilar login qiladi.', style='List Bullet')
    doc.add_paragraph('Admin modeli:', style='List Bullet')
    doc.add_paragraph('User bilan OneToOneField orqali bog‘langan.', style='List Bullet')
    doc.add_paragraph('AdminSerializer orqali yangi User (role="admin") va Admin obyekti yaratiladi.',
                      style='List Bullet')
    doc.add_paragraph('Update’da user va admin ma’lumotlari alohida yangilanadi.', style='List Bullet')
    doc.add_paragraph('Exploiter modeli:', style='List Bullet')
    doc.add_paragraph('User bilan OneToOneField orqali bog‘langan.', style='List Bullet')
    doc.add_paragraph('ExploiterSerializer orqali yangi User (role="exploiter") va Exploiter obyekti yaratiladi.',
                      style='List Bullet')
    doc.add_paragraph('Update’da user va exploiter ma’lumotlari alohida yangilanadi.', style='List Bullet')
    doc.add_paragraph('Ruxsatlar:', style='List Bullet')
    doc.add_paragraph('User va Admin faqat adminlar tomonidan boshqariladi (IsAdmin).', style='List Bullet')
    doc.add_paragraph(
        'Exploiter autentifikatsiya qilingan foydalanuvchilar tomonidan ko‘riladi (IsAuthenticatedOrExploiter).',
        style='List Bullet')

    # 9. Frontendchilar uchun Ko‘rsatmalar
    doc.add_heading('9. Frontendchilar uchun Ko‘rsatmalar', level=1)
    doc.add_paragraph(
        'Ushbu API frontend ishlab chiquvchilar uchun mo‘ljallangan bo‘lib, quyidagi ko‘rsatmalar bilan ishlaydi:')
    doc.add_paragraph('API Endpoint’lari:', style='List Bullet')
    doc.add_paragraph('Barcha endpoint’lar http://127.0.0.1:8000/app/v1/ prefiksi bilan ishlaydi.', style='List Bullet')
    doc.add_paragraph('Masalan: /app/v1/films/, /app/v1/comments/, /app/v1/ratings/.', style='List Bullet')
    doc.add_paragraph('JWT Autentifikatsiyasi:', style='List Bullet')
    doc.add_paragraph('POST /api/token/ orqali access va refresh tokenlari olinadi.', style='List Bullet')
    doc.add_paragraph('So‘rovlarda Authorization header’ida Bearer <access_token> ishlatiladi.', style='List Bullet')
    doc.add_paragraph('Access token muddati tugasa, POST /api/token/refresh/ orqali yangilanadi.', style='List Bullet')
    doc.add_paragraph('Swagger UI:', style='List Bullet')
    doc.add_paragraph('http://127.0.0.1:8000/swagger/ orqali barcha endpoint’larni sinab ko‘rish mumkin.',
                      style='List Bullet')
    doc.add_paragraph('“Authorize” tugmasi orqali Bearer token kiritiladi.', style='List Bullet')
    doc.add_paragraph('Maslahatlar:', style='List Bullet')
    doc.add_paragraph('Tokenlarni xavfsiz saqlang (masalan, HttpOnly cookie’larda).', style='List Bullet')
    doc.add_paragraph('So‘rovlar uchun Content-Type: application/json ishlatiladi.', style='List Bullet')
    doc.add_paragraph('Xatolarni tekshirish uchun Swagger’dagi response kodlarini ko‘ring (masalan, 401 Unauthorized).',
                      style='List Bullet')

    # Faylni saqlash
    doc.save('documentation.docx')


if __name__ == '__main__':
    create_documentation()